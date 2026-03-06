
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

MD_FILE = r"c:\Cladue ai project\DOCUMENTATION\Abstract.md"
OUTPUT_FILE = r"c:\Cladue ai project\DOCUMENTATION\Project_Abstract.docx"

def create_abstract():
    doc = Document()
    
    # Check file exists
    if not os.path.exists(MD_FILE):
        print("Abstract.md not found!")
        return

    with open(MD_FILE, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Title
        if line.startswith("# ABSTRACT"):
            heading = doc.add_heading("ABSTRACT", 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
            
        # Bold fields (Project Title, Domain, Keywords)
        # Simple parser for the specific format we wrote
        if line.startswith("**"):
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(12)
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(12)
            # Spacing
            p.paragraph_format.space_after = Pt(12)
            continue
            
        # Normal text
        p = doc.add_paragraph(line)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(12)
        run = p.runs[0]
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    doc.save(OUTPUT_FILE)
    print(f"Created {OUTPUT_FILE}")

if __name__ == "__main__":
    create_abstract()
