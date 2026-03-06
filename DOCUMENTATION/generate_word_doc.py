
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Configuration
DOC_DIR = r"c:\Cladue ai project\DOCUMENTATION"
FILES_TO_PROCESS = [
    "Part_1_Introduction_and_Survey.md",
    "Part_2_System_Analysis_and_Design.md",
    "Part_3_Implementation_and_Results.md"
]
OUTPUT_FILE = os.path.join(DOC_DIR, "Final_Year_Project_Report.docx")

def parse_markdown_line(doc, line):
    line = line.strip()
    if not line:
        return

    # 1. Headers
    match_header = re.match(r'^(#{1,6})\s+(.*)', line)
    if match_header:
        level = len(match_header.group(1))
        text = match_header.group(2)
        # Map MD headers to Docx logic (Title is massive, Heading 1 is chapter)
        if level == 1:
            doc.add_heading(text, 0) # Title style
        else:
            doc.add_heading(text, level)
        return

    # 2. Images: ![Alt](Path)
    # Regex handles standard markdown image syntax
    match_image = re.search(r'!\[(.*?)\]\((.*?)\)', line)
    if match_image:
        alt_text = match_image.group(1)
        image_path = match_image.group(2)
        
        # Clean path if it has "file:///" prefix or is relative
        if image_path.startswith("file:///"):
            image_path = image_path.replace("file:///", "")
        
        # Handle Windows paths issues if any
        image_path = image_path.replace("%20", " ")
        
        try:
            if os.path.exists(image_path):
                doc.add_picture(image_path, width=Inches(6.0))
                # Add caption
                para = doc.add_paragraph(alt_text)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.italic = True
            else:
                doc.add_paragraph(f"[Image Missing: {image_path}]", style='Body Text')
        except Exception as e:
            doc.add_paragraph(f"[Error loading image: {str(e)}]", style='Body Text')
        return

    # 3. Code Blocks (Very basic detection)
    if line.startswith("```"):
        return # Skip fence lines for now, essentially
    
    # 4. Bullet Points
    if line.startswith("- ") or line.startswith("* "):
        text = line[2:]
        p = doc.add_paragraph(text, style='List Bullet')
        apply_inline_formatting(p)
        return

    # 5. Numbered Lists
    match_num = re.match(r'^\d+\.\s+(.*)', line)
    if match_num:
        text = match_num.group(1)
        p = doc.add_paragraph(text, style='List Number')
        apply_inline_formatting(p)
        return
        
    # 6. Normal Text
    p = doc.add_paragraph(line)
    apply_inline_formatting(p)

def apply_inline_formatting(paragraph):
    # This is a basic formatter for **Bold** and *Italic* and `Code`
    # Note: This implementation replaces the whole paragraph logic for simplicity 
    # or would need complex run-splitting which is error prone in simple scripts.
    # For this version, we will just handle BOLD for specifically marked terms 
    # if they are at the start or straightforward.
    
    # A robust solution is complex, so we will do a simple pass:
    # If the text has basic markdown, we strip it to look clean at least.
    
    # Simple Bold parser for "**Text**"
    text = paragraph.text
    if "**" in text:
        # Clear existing runs
        paragraph.clear()
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

def create_document():
    doc = Document()
    
    # Add Title Page Content manually for better style
    doc.add_heading("ADVANCED AI CONTENT DETECTION SYSTEM", 0)
    p = doc.add_paragraph("Submitted in partial fulfillment of the requirements for the degree of Bachelor of Engineering")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    for file_name in FILES_TO_PROCESS:
        file_path = os.path.join(DOC_DIR, file_name)
        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            continue
            
        print(f"Processing {file_name}...")
        with open(file_path, "r", encoding="utf-8") as f:
            lines = f.readlines()
            
        for line in lines:
            parse_markdown_line(doc, line)
            
        # Add page break between chapters/parts
        doc.add_page_break()

    doc.save(OUTPUT_FILE)
    print(f"Successfully created: {OUTPUT_FILE}")

if __name__ == "__main__":
    create_document()
