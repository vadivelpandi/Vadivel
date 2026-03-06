
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configuration
DOC_DIR = r"c:\Cladue ai project\DOCUMENTATION"
SOURCE_FILE = "../SYSTEM_WORKFLOW.md"
OUTPUT_FILE = os.path.join(DOC_DIR, "workflow_generated.docx")

def set_cell_border(cell, **kwargs):
    """
    Helper to set cell borders for 'code block' look (if we used tables, 
    but for now we might just use shading or courier font).
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def add_code_block(doc, code_lines):
    """
    Adds a code block to the document using a table with a single cell
    to simulate a box.
    """
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    
    # Set background shading (light gray)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'E6E6E6') # Light gray
    tcPr.append(shd)
    
    paragraph = cell.paragraphs[0]
    paragraph.style = 'No Spacing'
    
    # Add text line by line
    for line in code_lines:
        run = paragraph.add_run(line + '\n')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        
def parse_markdown_line(doc, line):
    line = line.strip()
    if not line:
        return None

    # 1. Headers
    match_header = re.match(r'^(#{1,6})\s+(.*)', line)
    if match_header:
        level = len(match_header.group(1))
        text = match_header.group(2)
        if level == 1:
            doc.add_heading(text, 0) # Title
        else:
            doc.add_heading(text, level)
        return "header"

    # 2. Images: ![Alt](Path)
    match_image = re.search(r'!\[(.*?)\]\((.*?)\)', line)
    if match_image:
        alt_text = match_image.group(1)
        image_path = match_image.group(2)
        if image_path.startswith("file:///"):
            image_path = image_path.replace("file:///", "")
        image_path = image_path.replace("%20", " ")
        
        try:
            if os.path.exists(image_path):
                doc.add_picture(image_path, width=Inches(6.0))
                p = doc.add_paragraph(alt_text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.italic = True
            else:
                doc.add_paragraph(f"[Image: {alt_text}]", style='Body Text')
        except Exception as e:
            doc.add_paragraph(f"[Error loading image: {str(e)}]", style='Body Text')
        return "image"
        
    # 3. Code Blocks start/end
    if line.startswith("```"):
        return "code_fence"
    
    # 4. Bullet Points
    if line.startswith("- ") or line.startswith("* "):
        text = line[2:]
        p = doc.add_paragraph(text, style='List Bullet')
        apply_inline_formatting(p)
        return "list"

    # 5. Numbered Lists
    match_num = re.match(r'^\d+\.\s+(.*)', line)
    if match_num:
        text = match_num.group(1)
        p = doc.add_paragraph(text, style='List Number')
        apply_inline_formatting(p)
        return "list"
        
    # 6. Normal Text
    p = doc.add_paragraph(line)
    apply_inline_formatting(p)
    return "text"

def apply_inline_formatting(paragraph):
    text = paragraph.text
    # Simple Bold parser for "**Text**"
    if "**" in text:
        paragraph.clear()
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

def create_document():
    doc = Document()
    
    file_path = os.path.join(DOC_DIR, SOURCE_FILE)
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return

    print(f"Processing {SOURCE_FILE}...")
    with open(file_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    in_code_block = False
    code_buffer = []

    for line in lines:
        stripped = line.strip()
        
        # Handle Code Blocks specifically
        if stripped.startswith("```"):
            if in_code_block:
                # End of block
                add_code_block(doc, code_buffer)
                code_buffer = []
                in_code_block = False
            else:
                # Start of block
                in_code_block = True
            continue
            
        if in_code_block:
            code_buffer.append(line.rstrip()) # Keep indentation but remove newline for control
            continue
            
        parse_markdown_line(doc, line)

    doc.save(OUTPUT_FILE)
    print(f"Successfully created: {OUTPUT_FILE}")

if __name__ == "__main__":
    create_document()
